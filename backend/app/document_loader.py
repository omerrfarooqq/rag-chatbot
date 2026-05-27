import io
from typing import List
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from .config import CHUNK_SIZE, CHUNK_OVERLAP


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return _extract_txt(data)
    raise ValueError(f"Unsupported file type: {filename}")


def chunk_text(text: str, source: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [Document(page_content=c, metadata={"source": source}) for c in chunks]
